import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO 
import pandas as pd

#Configuracion de la Pagin
st.set_page_config(page_title="Generador de contenidos con OpenAI", layout="wide")
#Configuracion de OpenAI
cliente = OpenAI(api_key="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")

#Crear la funcion para generar articulos
def generar_articulo(topic):
    try:
        response = cliente.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role":"system","content":"Eres un extero en redaccion de articulos SEO"},
                {"role":"user","content":f"Escribe un articulo optimizado para SEO sobre: {topic}"}
            ],
            max_tokens = 1000
        )
        return response.choices[0].message.content.strip()   # Limpiando la respuestas
    except Exception as e:
        st.error(f"Error al generar el articulo: {str(e)}")
        return None
    
# Funcion para generar codigo
def generar_codigo(descripcion):
    try:
        response = cliente.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role":"system","content":"Eres un programador de python. Proporciona solo codigo python, sin explicaciones ni comentarios adicionales."},
                {"role":"user","content":f"Escribe el codigo de python para: {descripcion}. IMPORTANTE: proporciona SOLO el codigo, sin ningun texto explicativo."}
            ],
            temperature=0.2
        )
        code = response.choices[0].message.content.strip()   # Limpiando la respuestas
        
        # Elimina cualquier texto que no sea codigo
        code_lines = []
        in_code_block = False
        for line in code.split("\n"):
            if line.strip().startswith("```python"):
                in_code_block = True
                continue
            elif line.strip().startswith("```"):
                in_code_block = False
                continue
            elif in_code_block or not line.strip().startswith("```"):
                code_lines.append(line)
        return "\n".join(code_lines).strip()
    except Exception as e:
        st.error(f"Error al generar el codigo: {str(e)}")
        return None

#Crear la funcion para crear word
def generar_tabla(descripcion):
    
    try:
        response = cliente.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role":"system","content":"Genera datos en formato CSV con encabezados. Los datos deben ser coherentes y bien estructurados."},
                {"role":"user","content":f"Genera una tabla de datos en formato CSV para: {descripcion}. Incluye encabezados y al menos 10 filas de datos. Usa coma como separador."}
            ],
            temperature=0.5
        )
        return response.choices[0].message.content.strip()   # Limpiando la respuestas
    except Exception as e:
        st.error(f"Error al generar la tabla: {str(e)}")
        return None

#Crear la funcion para crear word
def create_word(article): 
    doc=Document()
    doc.add_heading("Articulo generado",0)
    doc.add_paragraph(article)
    buffer=BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# Crear la funcion para crear la Tabla
def create_excel(data):
    # Convertir el texto csv en dataframe
    lines = [line.strip() for line in data.split("\n") if line.strip()]
    if not lines:
        raise ValueError("No hay datos para procesar")
    # Separar encabezados de los datos
    headers=lines[2].split(",")
    data_filas=[line.split(",") for line in lines[2:-2]]
    
    # Crear Dataframe
    df = pd.DataFrame(data_filas, columns=headers)
    
    # Guardar en excel con formato mejorado
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Datos")
        # Ajustar el ancho de las columnas
        worksheet = writer.sheets["Datos"]
        for idx, col in enumerate(df.columns):   
            max_len = max(df[col].astype(str).apply(len).max(), len(col))
            worksheet.column_dimensions[chr(65 + idx)].width = max_len + 2

    buffer.seek(0)
    return buffer, df

#Titulo
st.title("Generador de Contenido con GPT-4o")
#Barra lateral
seccion = st.sidebar.selectbox("Selecciona:",["Generacion de Articulo","Generacion de codigos", "Generacion de tabla de datos"])


if seccion == "Generacion de Articulo":
    st.header("Generacion de Articulo")
    topic=st.text_input("Ingresar un tema para el articulo")
    if st.button("Generacion de Articulo"):
        if topic:
            with st.spinner("Generando...."):
                article=generar_articulo(topic)
                if article:
                    st.success("Articulo generado")
                    st.markdown("### Vista previa del articulo:")
                    st.markdown(article)
                    
                    # boton de descarga
                    st.download_button(
                        "Descargar como Word",
                        data=create_word(article),
                        file_name="articulo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )  
        else:
            st.warning("Ingrese un tema:")              
elif seccion == "Generacion de codigos":
    st.header("Generacion de Codigo")
    descripcion=st.text_input("Cual es el codigo que necesitas")
    if st.button("Generacion de Articulo"):
        if descripcion:
            with st.spinner("Generando...."):
                code=generar_codigo(descripcion)
                if code:
                    st.success("Codigo generado exitosamente")
                    st.code(code,language="python")
                    # boton de descarga
                    st.download_button(
                        "Descargar como Python",
                        data=code,
                        file_name="codigo.py",
                        mime="text/plain"
                    )
        else:
            st.warning("Describe el codigo que necesitas")            
elif seccion == "Generacion de tabla de datos":
    st.header("Generacion de Tabla")
    data_descripcion=st.text_input("Describe la tabla de datos que necesitas")    
# Boton para generar 
    if st.button("Generar tabla"):
        if data_descripcion:
            with st.spinner("Generando tabla de datos...."):
                data=generar_tabla(data_descripcion)
                
                if data:
                    st.success("Tabla generada exitosamente")
                    # Crear Excel y obtener dataframe

                    excel_file, df=create_excel(data)
                    
                    if excel_file and df is not None:
                        #Mostrar vista previa
                        st.markdown("### Vista previa de la tabla")
                        st.dataframe(df, use_container_width=True)
                    # boton de descarga
                        st.download_button(
                        "Descargar como Excel",
                        data=excel_file,
                        file_name="tabla_datos.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
        else:
            st.warning("Describe la tabla que necesitas")
  







